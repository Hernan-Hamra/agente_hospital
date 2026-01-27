"""
Script para convertir documentos PDF/DOCX a JSON flat (sin estructura jerárquica)
Enfoque simple: chunks de tamaño fijo con overlap, metadata clara por chunk

Inspirado en el método de Patricia: dividir documentos y agregar metadata en cada pedazo
Mejoras: automático, con tablas detectadas, overlap entre chunks
"""
import os
import json
from pathlib import Path
from typing import Dict, List, Tuple
import pdfplumber
from docx import Document
import re


# Regex para detectar contactos en texto
PHONE_PATTERNS = [
    r'\b0810[-\s]?\d{3}[-\s]?\d{4}\b',  # 0810-xxx-xxxx
    r'\b\d{2}[-\s]?\d{8}\b',              # 11-66075765
    r'\b\d{4}[-\s]?\d{4}\b',              # 4716-8723
    r'\b\d{4}[-\s]?\d{4}\s*(?:int|Int|INT)[.:]\s*\d+(?:/\d+)*\b',  # 4716-8723 int. 347/381
]

EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'


class DocumentToJsonFlat:
    """Convierte documentos a JSON flat con chunks uniformes"""

    def __init__(self, data_path: str, output_path: str, chunk_size: int = 1800, overlap: int = 300):
        """
        Args:
            data_path: Ruta a documentos originales
            output_path: Ruta de salida para JSONs
            chunk_size: Tamaño de cada chunk en caracteres
            overlap: Superposición entre chunks consecutivos
        """
        self.data_path = Path(data_path)
        self.output_path = Path(output_path)
        self.output_path.mkdir(parents=True, exist_ok=True)
        self.chunk_size = chunk_size
        self.overlap = overlap

    def extract_text_and_tables_from_pdf(self, pdf_path: Path) -> tuple:
        """
        Extrae texto completo y tablas de PDF

        Returns:
            (texto_completo: str, tablas: list)
        """
        texto_partes = []
        tablas = []

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extraer TABLAS
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    tabla_data = {
                        "numero": len(tablas) + 1,
                        "pagina": page_num,
                        "filas": len(table),
                        "columnas": len(table[0]) if table else 0,
                        "contenido": table
                    }
                    tablas.append(tabla_data)

                # Extraer TEXTO
                page_text = page.extract_text()
                if page_text:
                    texto_partes.append(page_text.strip())

        texto_completo = "\n\n".join(texto_partes)
        return texto_completo, tablas

    def extract_text_and_tables_from_docx(self, docx_path: Path) -> tuple:
        """
        Extrae texto completo y tablas de DOCX

        Returns:
            (texto_completo: str, tablas: list)
        """
        doc = Document(docx_path)

        # Extraer TABLAS primero (son críticas)
        tablas = []
        for table_idx, table in enumerate(doc.tables):
            tabla_data = {
                "numero": table_idx + 1,
                "filas": len(table.rows),
                "columnas": len(table.columns),
                "contenido": []
            }

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                tabla_data["contenido"].append(row_data)

            tablas.append(tabla_data)

        # Extraer TEXTO completo
        texto_parrafos = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texto_parrafos.append(text)

        texto_completo = "\n".join(texto_parrafos)

        return texto_completo, tablas

    def create_chunks_with_overlap(self, texto: str, metadata_base: dict) -> List[Dict]:
        """
        Divide texto en chunks de tamaño fijo con overlap

        Args:
            texto: Texto completo a dividir
            metadata_base: Metadata común (obra_social, archivo)

        Returns:
            Lista de chunks con metadata
        """
        chunks = []

        # Si el texto es muy corto, un solo chunk
        if len(texto) <= self.chunk_size:
            chunks.append({
                **metadata_base,
                "chunk_id": "001",
                "texto": texto,
                "es_tabla": False
            })
            return chunks

        # Dividir con overlap
        start = 0
        chunk_num = 1

        while start < len(texto):
            # Extraer chunk
            end = start + self.chunk_size
            chunk_text = texto[start:end]

            # Intentar cortar en un espacio/salto de línea para no partir palabras
            if end < len(texto):
                # Buscar último espacio o salto de línea en los últimos 100 chars
                last_space = chunk_text.rfind(' ', -100)
                last_newline = chunk_text.rfind('\n', -100)
                cut_point = max(last_space, last_newline)

                if cut_point > 0:
                    chunk_text = chunk_text[:cut_point]
                    end = start + cut_point

            # Crear chunk
            chunks.append({
                **metadata_base,
                "chunk_id": f"{chunk_num:03d}",
                "texto": chunk_text.strip(),
                "es_tabla": False
            })

            chunk_num += 1

            # Avanzar con overlap
            start = end - self.overlap

            # Evitar loop infinito
            if start + self.chunk_size >= len(texto) and start < len(texto):
                # Último chunk
                chunk_text = texto[start:]
                if chunk_text.strip():
                    chunks.append({
                        **metadata_base,
                        "chunk_id": f"{chunk_num:03d}",
                        "texto": chunk_text.strip(),
                        "es_tabla": False
                    })
                break

        return chunks

    def is_contact_table(self, tabla: dict) -> bool:
        """
        Detecta si es una tabla de contactos (mail, teléfono).
        Mira los encabezados (primera fila) para decidir.
        """
        if not tabla.get("contenido") or len(tabla["contenido"]) < 2:
            return False

        encabezados = tabla["contenido"][0]
        encabezados_lower = [str(h).lower() for h in encabezados if h]

        # Palabras clave que indican tabla de contactos
        contact_keywords = ["mail", "tel", "teléfono", "telefono", "email", "correo", "contacto"]

        for header in encabezados_lower:
            for keyword in contact_keywords:
                if keyword in header:
                    return True
        return False

    def contact_table_to_sentences(self, tabla: dict, obra_social: str) -> str:
        """
        Convierte tabla de contactos a oraciones para mejor RAG.

        Ejemplo:
        | Sector | Tel | mail |
        | Mesa Operativa | 0810-888 | auto@asi.com |

        →
        El mail de Mesa Operativa de ASI es auto@asi.com.
        El teléfono de Mesa Operativa de ASI es 0810-888.
        """
        contenido = tabla.get("contenido", [])
        if len(contenido) < 2:
            return ""

        encabezados = contenido[0]
        oraciones = []

        # Encontrar índices de columnas relevantes
        idx_nombre = None  # Columna con el nombre (Sector, Área, etc.)
        idx_mail = None
        idx_tel = None

        for i, header in enumerate(encabezados):
            header_lower = str(header).lower() if header else ""
            if any(kw in header_lower for kw in ["sector", "área", "area", "nombre", "departamento"]):
                idx_nombre = i
            elif any(kw in header_lower for kw in ["mail", "email", "correo"]):
                idx_mail = i
            elif any(kw in header_lower for kw in ["tel", "teléfono", "telefono"]):
                idx_tel = i

        # Si no encontró columna de nombre, usar la primera
        if idx_nombre is None:
            idx_nombre = 0

        # Procesar filas (saltar encabezado)
        for row in contenido[1:]:
            if not row or len(row) <= idx_nombre:
                continue

            nombre = str(row[idx_nombre]).strip() if row[idx_nombre] else ""
            if not nombre:
                continue

            # Generar oraciones para mail
            if idx_mail is not None and len(row) > idx_mail and row[idx_mail]:
                mail_value = str(row[idx_mail]).strip()
                if mail_value and "@" in mail_value:
                    oraciones.append(f"El mail de {nombre} de {obra_social} es {mail_value}.")

            # Generar oraciones para teléfono
            if idx_tel is not None and len(row) > idx_tel and row[idx_tel]:
                tel_value = str(row[idx_tel]).strip()
                if tel_value:
                    oraciones.append(f"El teléfono de {nombre} de {obra_social} es {tel_value}.")

        return "\n".join(oraciones)

    def _detect_table_keywords(self, tabla: dict) -> str:
        """
        Detecta el tipo de tabla y genera keywords semánticas para mejorar RAG.
        Estas keywords ayudan al embedding a encontrar la tabla con queries coloquiales.
        """
        contenido = tabla.get("contenido", [])
        contenido_str = str(contenido).lower()

        keywords = []

        # Detectar tablas de coseguros/valores/precios
        if "$" in contenido_str or "coseguro" in contenido_str or "valor" in contenido_str:
            keywords.extend(["COSEGUROS", "VALORES", "PRECIOS", "TARIFAS", "CUANTO CUESTA", "CUANTO SALE"])

        # Detectar consultas médicas
        if "médico" in contenido_str or "medico" in contenido_str or "especialista" in contenido_str:
            keywords.extend(["CONSULTA", "ESPECIALISTA", "MEDICO", "ATENCION"])

        # Detectar laboratorio
        if "laboratorio" in contenido_str or "determinacion" in contenido_str:
            keywords.extend(["LABORATORIO", "ANALISIS", "ESTUDIOS"])

        # Detectar imágenes/radiología
        if "imagen" in contenido_str or "eco" in contenido_str or "tac" in contenido_str or "rmn" in contenido_str:
            keywords.extend(["IMAGENES", "RADIOLOGIA", "ECOGRAFIA", "TOMOGRAFIA", "RESONANCIA"])

        # Detectar kinesiología/rehabilitación
        if "kinesio" in contenido_str or "rehabilitacion" in contenido_str or "fisio" in contenido_str:
            keywords.extend(["KINESIOLOGIA", "REHABILITACION", "FISIOTERAPIA", "SESION"])

        # Detectar cobertura/autorización
        if "cobertura" in contenido_str or "autorizacion" in contenido_str or "autorización" in contenido_str:
            keywords.extend(["COBERTURA", "AUTORIZACION", "REQUIERE", "NECESITA"])

        # Detectar internación
        if "internacion" in contenido_str or "internación" in contenido_str:
            keywords.extend(["INTERNACION", "HOSPITALIZACION", "CIRUGIA"])

        # Detectar planes
        if "plan" in contenido_str or "delta" in contenido_str or "quantum" in contenido_str:
            keywords.extend(["PLANES", "TIPO DE PLAN", "CATEGORIA"])

        # Detectar prácticas/códigos
        if any(c.isdigit() for c in contenido_str[:200]):  # Códigos numéricos
            keywords.extend(["PRACTICAS", "CODIGOS", "NOMENCLADOR"])

        return " ".join(keywords) if keywords else ""

    def table_to_text(self, tabla: dict, obra_social: str = "") -> str:
        """
        Convierte tabla a formato texto.
        - Tablas de contacto → oraciones (mejor para RAG)
        - Otras tablas → formato tabla con keywords semánticas (mejor para búsqueda)
        """
        # Si es tabla de contactos, convertir a oraciones
        if self.is_contact_table(tabla):
            return self.contact_table_to_sentences(tabla, obra_social)

        # Detectar keywords semánticas para mejorar retrieval
        keywords = self._detect_table_keywords(tabla)
        keyword_header = f"{keywords}\n" if keywords else ""

        # Formato tabla normal con keywords al inicio
        texto_tabla = f"{keyword_header}TABLA #{tabla['numero']} ({tabla['filas']} filas x {tabla['columnas']} columnas)\n\n"

        for row_idx, row in enumerate(tabla["contenido"]):
            # Primera fila como encabezado
            if row_idx == 0:
                texto_tabla += " | ".join([f"**{cell}**" for cell in row if cell]) + "\n"
                texto_tabla += "---\n"
            else:
                texto_tabla += " | ".join([f"{cell}" for cell in row if cell]) + "\n"

        return texto_tabla.strip()

    def extract_contacts_from_text(self, texto: str, obra_social: str) -> Tuple[str, List[str]]:
        """
        Extrae teléfonos y emails del texto y genera oraciones de contacto.

        Args:
            texto: Texto completo del documento
            obra_social: Nombre de la obra social

        Returns:
            (texto_contactos: str, contactos_encontrados: list)
        """
        contactos = []
        emails_encontrados = []
        telefonos_encontrados = []

        # Buscar emails
        emails = re.findall(EMAIL_PATTERN, texto, re.IGNORECASE)
        for email in emails:
            if email not in emails_encontrados:
                emails_encontrados.append(email)

        # Buscar teléfonos
        for pattern in PHONE_PATTERNS:
            telefonos = re.findall(pattern, texto)
            for tel in telefonos:
                if tel not in telefonos_encontrados:
                    telefonos_encontrados.append(tel)

        # Generar oraciones de contacto
        # Intentar extraer contexto (qué área/sector corresponde a cada contacto)
        lineas = texto.split('\n')

        for email in emails_encontrados:
            # Buscar contexto cercano al email
            contexto = self._find_contact_context(lineas, email)
            if contexto:
                contactos.append(f"El mail de {contexto} de {obra_social} es {email}.")
            else:
                contactos.append(f"El mail de contacto de {obra_social} es {email}.")

        for tel in telefonos_encontrados:
            # Buscar contexto cercano al teléfono
            contexto = self._find_contact_context(lineas, tel)
            if contexto:
                contactos.append(f"El teléfono de {contexto} de {obra_social} es {tel}.")
            else:
                contactos.append(f"El teléfono de contacto de {obra_social} es {tel}.")

        texto_contactos = "\n".join(contactos) if contactos else ""
        return texto_contactos, emails_encontrados + telefonos_encontrados

    def _find_contact_context(self, lineas: List[str], contacto: str) -> str:
        """
        Busca el contexto (área/sector) de un contacto en las líneas cercanas.

        Args:
            lineas: Lista de líneas del texto
            contacto: Email o teléfono a buscar

        Returns:
            Contexto encontrado o string vacío
        """
        for i, linea in enumerate(lineas):
            if contacto in linea:
                # Buscar en la misma línea antes del contacto
                # Ej: "Mesa Operativa: autorizaciones@asi.com.ar"
                partes = linea.split(contacto)
                if partes[0].strip():
                    # Limpiar y extraer nombre del área
                    contexto = partes[0].strip()
                    # Quitar caracteres de puntuación al final
                    contexto = re.sub(r'[:\-–•\s]+$', '', contexto)
                    # Tomar solo las últimas palabras relevantes
                    palabras = contexto.split()
                    if len(palabras) > 0:
                        # Tomar últimas 3-4 palabras como máximo
                        return ' '.join(palabras[-4:])

                # Buscar en línea anterior
                if i > 0:
                    linea_anterior = lineas[i-1].strip()
                    if linea_anterior and len(linea_anterior) < 50:
                        return linea_anterior

        return ""

    def process_document(self, file_path: Path, obra_social: str) -> List[Dict]:
        """
        Procesa un documento y genera chunks flat

        Args:
            file_path: Ruta al archivo
            obra_social: Nombre de la obra social

        Returns:
            Lista de chunks con metadata
        """
        print(f"  Procesando: {file_path.name}")

        metadata_base = {
            "obra_social": obra_social.upper(),
            "archivo": file_path.name
        }

        all_chunks = []

        if file_path.suffix.lower() == '.pdf':
            # Procesar PDF (texto + tablas)
            texto_completo, tablas = self.extract_text_and_tables_from_pdf(file_path)

            # 1. Crear chunk de CONTACTOS extraídos del texto (PRIMERO, alta prioridad)
            texto_contactos, contactos_list = self.extract_contacts_from_text(texto_completo, obra_social.upper())
            if texto_contactos:
                all_chunks.append({
                    **metadata_base,
                    "chunk_id": "CONTACTOS",
                    "texto": texto_contactos,
                    "es_tabla": False,
                    "es_contactos": True,
                    "contactos_extraidos": contactos_list
                })
                print(f"    📞 Contactos extraídos: {len(contactos_list)}")

            # 2. Crear chunks de TABLAS (completas, sin dividir)
            for tabla in tablas:
                tabla_texto = self.table_to_text(tabla, obra_social.upper())
                all_chunks.append({
                    **metadata_base,
                    "chunk_id": f"T{tabla['numero']:03d}",
                    "texto": tabla_texto,
                    "es_tabla": True,
                    "tabla_numero": tabla['numero'],
                    "pagina": tabla['pagina']
                })

            # 3. Crear chunks de TEXTO con overlap
            chunks_texto = self.create_chunks_with_overlap(texto_completo, metadata_base)
            all_chunks.extend(chunks_texto)

        elif file_path.suffix.lower() == '.docx':
            # Procesar DOCX (texto + tablas)
            texto_completo, tablas = self.extract_text_and_tables_from_docx(file_path)

            # 1. Crear chunk de CONTACTOS extraídos del texto (PRIMERO, alta prioridad)
            texto_contactos, contactos_list = self.extract_contacts_from_text(texto_completo, obra_social.upper())
            if texto_contactos:
                all_chunks.append({
                    **metadata_base,
                    "chunk_id": "CONTACTOS",
                    "texto": texto_contactos,
                    "es_tabla": False,
                    "es_contactos": True,
                    "contactos_extraidos": contactos_list
                })
                print(f"    📞 Contactos extraídos: {len(contactos_list)}")

            # 2. Crear chunks de TABLAS (completas, sin dividir)
            for tabla in tablas:
                tabla_texto = self.table_to_text(tabla, obra_social.upper())
                all_chunks.append({
                    **metadata_base,
                    "chunk_id": f"T{tabla['numero']:03d}",
                    "texto": tabla_texto,
                    "es_tabla": True,
                    "tabla_numero": tabla['numero']
                })

            # 3. Crear chunks de TEXTO con overlap
            chunks_texto = self.create_chunks_with_overlap(texto_completo, metadata_base)
            all_chunks.extend(chunks_texto)

        print(f"    → {len(all_chunks)} chunks generados")
        return all_chunks

    def process_all_documents(self) -> dict:
        """
        Procesa todos los documentos de obras sociales

        Returns:
            Estadísticas del proceso
        """
        print("\n" + "="*70)
        print("CONVERSIÓN A JSON FLAT (chunks uniformes con overlap)")
        print("="*70)
        print(f"Chunk size: {self.chunk_size} chars")
        print(f"Overlap: {self.overlap} chars")
        print("="*70 + "\n")

        stats = {
            "obras_sociales": 0,
            "documentos": 0,
            "chunks_totales": 0,
            "chunks_tablas": 0
        }

        # Procesar cada obra social
        for os_dir in self.data_path.iterdir():
            if not os_dir.is_dir():
                continue

            obra_social = os_dir.name
            print(f"\n{'─'*70}")
            print(f"Obra Social: {obra_social.upper()}")
            print(f"{'─'*70}")

            os_output_dir = self.output_path / obra_social
            os_output_dir.mkdir(exist_ok=True)

            # Procesar cada archivo
            for file_path in os_dir.iterdir():
                if file_path.suffix.lower() not in ['.pdf', '.docx']:
                    continue

                # Generar chunks
                chunks = self.process_document(file_path, obra_social)

                # Contar tablas
                chunks_tablas = sum(1 for c in chunks if c.get("es_tabla", False))

                # Guardar JSON
                output_file = os_output_dir / f"{file_path.stem}_chunks_flat.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, ensure_ascii=False, indent=2)

                print(f"    ✅ Guardado: {output_file.name}")
                print(f"       Chunks texto: {len(chunks) - chunks_tablas}")
                print(f"       Chunks tablas: {chunks_tablas}")

                stats["documentos"] += 1
                stats["chunks_totales"] += len(chunks)
                stats["chunks_tablas"] += chunks_tablas

            stats["obras_sociales"] += 1

        # Resumen
        print("\n" + "="*70)
        print("RESUMEN")
        print("="*70)
        print(f"✅ Obras sociales: {stats['obras_sociales']}")
        print(f"✅ Documentos procesados: {stats['documentos']}")
        print(f"✅ Total chunks: {stats['chunks_totales']}")
        print(f"   - Chunks de texto: {stats['chunks_totales'] - stats['chunks_tablas']}")
        print(f"   - Chunks de tablas: {stats['chunks_tablas']}")
        print(f"\n📁 JSONs guardados en: {self.output_path}")
        print("="*70 + "\n")

        return stats


def main():
    # Configuración
    # Rutas relativas desde la raíz del proyecto
    project_root = Path(__file__).parent.parent.parent
    data_path = project_root / "data" / "obras_sociales"
    output_path = project_root / "data" / "obras_sociales_json"

    # Parámetros optimizados para bge-large-en-v1.5
    chunk_size = 1800  # ~450 tokens (de 512 max)
    overlap = 300      # ~75 tokens de overlap

    # Convertir obras sociales
    converter = DocumentToJsonFlat(
        data_path=str(data_path),
        output_path=str(output_path),
        chunk_size=chunk_size,
        overlap=overlap
    )

    stats = converter.process_all_documents()

    # También procesar grupo_pediatrico (está en carpeta separada)
    grupo_ped_path = project_root / "data" / "grupo_pediatrico"
    if grupo_ped_path.exists():
        print("\n" + "="*70)
        print("PROCESANDO GRUPO PEDIATRICO")
        print("="*70)

        # Crear salida para grupo_pediatrico
        grupo_ped_output = output_path / "grupo_pediatrico"
        grupo_ped_output.mkdir(exist_ok=True)

        for file_path in grupo_ped_path.iterdir():
            if file_path.suffix.lower() in ['.pdf', '.docx']:
                chunks = converter.process_document(file_path, "GRUPO_PEDIATRICO")

                # Guardar JSON
                output_file = grupo_ped_output / f"{file_path.stem}_chunks_flat.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, ensure_ascii=False, indent=2)

                print(f"    ✅ Guardado: {output_file.name}")
                stats["documentos"] += 1
                stats["chunks_totales"] += len(chunks)

    # Mostrar ejemplo de chunk
    print("\n" + "="*70)
    print("EJEMPLO DE CHUNK")
    print("="*70)

    # Buscar primer chunk de ASI
    example_file = Path(output_path) / "asi" / "2024-01-04_normas_chunks_flat.json"
    if example_file.exists():
        with open(example_file, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
            if chunks:
                print("\nPrimer chunk de ASI:")
                print(json.dumps(chunks[0], ensure_ascii=False, indent=2))

    return stats


if __name__ == "__main__":
    main()
