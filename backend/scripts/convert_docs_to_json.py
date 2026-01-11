"""
Script para convertir documentos PDF/DOCX a JSON estructurado
Cada JSON tendrá metadata clara de la obra social para evitar mezclas
"""
import os
import json
from pathlib import Path
from typing import Dict, List
import pdfplumber
from docx import Document
import re


class DocumentToJsonConverter:
    """Convierte documentos a JSON estructurado por obra social"""

    def __init__(self, data_path: str, output_path: str):
        self.data_path = Path(data_path)
        self.output_path = Path(output_path)
        self.output_path.mkdir(parents=True, exist_ok=True)

    def extract_pdf_content(self, pdf_path: Path, obra_social: str) -> Dict:
        """
        Extrae contenido de PDF manteniendo estructura

        Returns:
            Dict con metadata y contenido estructurado
        """
        print(f"  Procesando PDF: {pdf_path.name}")

        content = {
            "obra_social": obra_social.upper(),
            "archivo_original": pdf_path.name,
            "tipo": "pdf",
            "paginas": []
        }

        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    # Limpiar texto
                    page_text = page_text.strip()

                    content["paginas"].append({
                        "numero": i + 1,
                        "texto": page_text,
                        "obra_social": obra_social.upper()  # Redundante pero importante
                    })

        return content

    def extract_docx_content(self, docx_path: Path, obra_social: str) -> Dict:
        """
        Extrae contenido de DOCX manteniendo estructura (párrafos Y tablas)

        Returns:
            Dict con metadata y contenido estructurado
        """
        print(f"  Procesando DOCX: {docx_path.name}")

        content = {
            "obra_social": obra_social.upper(),
            "archivo_original": docx_path.name,
            "tipo": "docx",
            "secciones": [],
            "tablas": []
        }

        doc = Document(docx_path)

        # 1. Extraer TABLAS (importante para copagos, valores, etc.)
        print(f"    → Encontradas {len(doc.tables)} tablas")
        for table_idx, table in enumerate(doc.tables):
            tabla_data = {
                "numero": table_idx + 1,
                "filas": len(table.rows),
                "columnas": len(table.columns),
                "obra_social": obra_social.upper(),
                "contenido": []
            }

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                tabla_data["contenido"].append(row_data)

            content["tablas"].append(tabla_data)

        # 2. Agrupar párrafos en secciones (separados por títulos o saltos)
        current_section = {
            "titulo": "Introducción",
            "parrafos": [],
            "obra_social": obra_social.upper()
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Si parece un título (todo mayúsculas, corto, o bold)
            is_title = (
                text.isupper() and len(text) < 100
            ) or (
                para.style.name.startswith('Heading')
            )

            if is_title and current_section["parrafos"]:
                # Guardar sección actual y empezar nueva
                content["secciones"].append(current_section)
                current_section = {
                    "titulo": text,
                    "parrafos": [],
                    "obra_social": obra_social.upper()
                }
            else:
                current_section["parrafos"].append(text)

        # Agregar última sección
        if current_section["parrafos"]:
            content["secciones"].append(current_section)

        return content

    def create_chunks_from_json(self, content: Dict, chunk_size: int = 1000) -> List[Dict]:
        """
        Crea chunks desde el JSON estructurado
        Cada chunk mantiene la metadata de obra social
        """
        chunks = []
        obra_social = content["obra_social"]
        archivo = content["archivo_original"]

        if content["tipo"] == "pdf":
            # Procesar páginas
            for page in content["paginas"]:
                texto = page["texto"]
                page_num = page["numero"]

                # Si la página es muy larga, dividirla
                if len(texto) > chunk_size:
                    # Dividir por párrafos primero
                    parrafos = texto.split('\n')
                    current_chunk = ""

                    for parrafo in parrafos:
                        if len(current_chunk) + len(parrafo) > chunk_size and current_chunk:
                            chunks.append({
                                "obra_social": obra_social,
                                "archivo": archivo,
                                "pagina": page_num,
                                "texto": current_chunk.strip(),
                                "tipo": "pdf"
                            })
                            current_chunk = parrafo
                        else:
                            current_chunk += "\n" + parrafo

                    if current_chunk.strip():
                        chunks.append({
                            "obra_social": obra_social,
                            "archivo": archivo,
                            "pagina": page_num,
                            "texto": current_chunk.strip(),
                            "tipo": "pdf"
                        })
                else:
                    chunks.append({
                        "obra_social": obra_social,
                        "archivo": archivo,
                        "pagina": page_num,
                        "texto": texto,
                        "tipo": "pdf"
                    })

        elif content["tipo"] == "docx":
            # 1. Primero procesar TABLAS (muy importante para copagos, valores)
            for tabla in content.get("tablas", []):
                # Convertir tabla a texto estructurado
                tabla_texto = f"TABLA #{tabla['numero']} ({tabla['filas']} filas x {tabla['columnas']} columnas)\n\n"

                for row_idx, row in enumerate(tabla["contenido"]):
                    # Formato: Col1: valor | Col2: valor | Col3: valor
                    row_str = " | ".join([f"{cell}" for cell in row if cell])
                    tabla_texto += f"{row_str}\n"

                # Crear un chunk por cada tabla (las tablas son info crítica)
                chunks.append({
                    "obra_social": obra_social,
                    "archivo": archivo,
                    "seccion": f"Tabla {tabla['numero']}",
                    "texto": tabla_texto.strip(),
                    "tipo": "docx-tabla",
                    "es_tabla": True
                })

            # 2. Luego procesar secciones de texto
            for seccion in content["secciones"]:
                titulo = seccion["titulo"]
                texto_completo = f"{titulo}\n\n" + "\n".join(seccion["parrafos"])

                if len(texto_completo) > chunk_size:
                    # Dividir en chunks más pequeños
                    parrafos = seccion["parrafos"]
                    current_chunk = f"{titulo}\n\n"

                    for parrafo in parrafos:
                        if len(current_chunk) + len(parrafo) > chunk_size and len(current_chunk) > len(titulo) + 10:
                            chunks.append({
                                "obra_social": obra_social,
                                "archivo": archivo,
                                "seccion": titulo,
                                "texto": current_chunk.strip(),
                                "tipo": "docx"
                            })
                            current_chunk = f"{titulo} (continuación)\n\n" + parrafo
                        else:
                            current_chunk += "\n" + parrafo

                    if current_chunk.strip():
                        chunks.append({
                            "obra_social": obra_social,
                            "archivo": archivo,
                            "seccion": titulo,
                            "texto": current_chunk.strip(),
                            "tipo": "docx"
                        })
                else:
                    chunks.append({
                        "obra_social": obra_social,
                        "archivo": archivo,
                        "seccion": titulo,
                        "texto": texto_completo,
                        "tipo": "docx"
                    })

        return chunks

    def validate_chunks(self, chunks: List[Dict], obra_social: str) -> bool:
        """
        Valida que todos los chunks tengan el nombre de obra social correcto
        """
        print(f"\n  Validando {len(chunks)} chunks...")
        errors = []

        for i, chunk in enumerate(chunks):
            # Verificar que tenga obra_social
            if "obra_social" not in chunk:
                errors.append(f"Chunk {i}: falta campo 'obra_social'")
                continue

            # Verificar que sea la correcta
            if chunk["obra_social"] != obra_social.upper():
                errors.append(f"Chunk {i}: obra_social incorrecta (esperada: {obra_social.upper()}, encontrada: {chunk['obra_social']})")

            # Verificar que tenga texto
            if not chunk.get("texto", "").strip():
                errors.append(f"Chunk {i}: texto vacío")

            # Opcional: verificar que el texto mencione la obra social
            texto_lower = chunk.get("texto", "").lower()
            obra_social_lower = obra_social.lower()

            # Buscar menciones de la obra social en el texto
            # (esto es informativo, no un error crítico)
            if obra_social_lower in texto_lower or obra_social.upper() in chunk.get("texto", ""):
                # Bien, el texto menciona la obra social
                pass

        if errors:
            print(f"  ⚠️  Se encontraron {len(errors)} errores:")
            for error in errors[:5]:  # Mostrar solo primeros 5
                print(f"    - {error}")
            return False
        else:
            print(f"  ✅ Todos los chunks son válidos")
            return True

    def process_all_documents(self, docs_general_path: str = None):
        """
        Procesa todos los documentos y genera JSONs estructurados

        Args:
            docs_general_path: Ruta opcional a documentos generales del hospital
        """
        print("\n" + "="*60)
        print("CONVERSIÓN DE DOCUMENTOS A JSON ESTRUCTURADO")
        print("="*60 + "\n")

        all_stats = {
            "obras_sociales_procesadas": 0,
            "documentos_procesados": 0,
            "documentos_generales": 0,
            "chunks_totales": 0,
            "validacion_exitosa": True
        }

        # 1. Procesar cada obra social
        for os_dir in self.data_path.iterdir():
            if not os_dir.is_dir():
                continue

            obra_social = os_dir.name
            print(f"\n{'─'*60}")
            print(f"Obra Social: {obra_social.upper()}")
            print(f"{'─'*60}")

            os_output_dir = self.output_path / obra_social
            os_output_dir.mkdir(exist_ok=True)

            # Procesar cada archivo
            for file_path in os_dir.iterdir():
                if file_path.suffix.lower() == '.pdf':
                    content = self.extract_pdf_content(file_path, obra_social)
                elif file_path.suffix.lower() == '.docx':
                    content = self.extract_docx_content(file_path, obra_social)
                else:
                    continue

                # Guardar JSON completo (para referencia)
                json_path = os_output_dir / f"{file_path.stem}_estructura.json"
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(content, f, ensure_ascii=False, indent=2)
                print(f"    ✅ JSON estructura: {json_path.name}")

                # Crear chunks
                chunks = self.create_chunks_from_json(content, chunk_size=1000)
                print(f"    📦 Generados {len(chunks)} chunks")

                # Validar chunks
                is_valid = self.validate_chunks(chunks, obra_social)
                if not is_valid:
                    all_stats["validacion_exitosa"] = False

                # Guardar chunks
                chunks_path = os_output_dir / f"{file_path.stem}_chunks.json"
                with open(chunks_path, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, ensure_ascii=False, indent=2)
                print(f"    ✅ JSON chunks: {chunks_path.name}")

                all_stats["documentos_procesados"] += 1
                all_stats["chunks_totales"] += len(chunks)

            all_stats["obras_sociales_procesadas"] += 1

        # 2. Procesar documentos generales del hospital (checklist, protocolos, etc.)
        if docs_general_path:
            docs_path = Path(docs_general_path)
            if docs_path.exists():
                print(f"\n{'─'*60}")
                print(f"Documentos GENERALES del Hospital")
                print(f"{'─'*60}")

                general_output_dir = self.output_path / "general"
                general_output_dir.mkdir(exist_ok=True)

                for file_path in docs_path.glob("*.docx"):
                    print(f"  Procesando: {file_path.name}")
                    content = self.extract_docx_content(file_path, "GENERAL")

                    # Guardar JSON completo
                    json_path = general_output_dir / f"{file_path.stem}_estructura.json"
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(content, f, ensure_ascii=False, indent=2)
                    print(f"    ✅ JSON estructura: {json_path.name}")

                    # Crear chunks
                    chunks = self.create_chunks_from_json(content, chunk_size=1000)
                    print(f"    📦 Generados {len(chunks)} chunks")

                    # Validar chunks
                    is_valid = self.validate_chunks(chunks, "GENERAL")
                    if not is_valid:
                        all_stats["validacion_exitosa"] = False

                    # Guardar chunks
                    chunks_path = general_output_dir / f"{file_path.stem}_chunks.json"
                    with open(chunks_path, 'w', encoding='utf-8') as f:
                        json.dump(chunks, f, ensure_ascii=False, indent=2)
                    print(f"    ✅ JSON chunks: {chunks_path.name}")

                    all_stats["documentos_generales"] += 1
                    all_stats["documentos_procesados"] += 1
                    all_stats["chunks_totales"] += len(chunks)

        # Resumen final
        print("\n" + "="*60)
        print("RESUMEN DE CONVERSIÓN")
        print("="*60)
        print(f"✅ Obras sociales procesadas: {all_stats['obras_sociales_procesadas']}")
        print(f"✅ Documentos de obras sociales: {all_stats['documentos_procesados'] - all_stats['documentos_generales']}")
        print(f"✅ Documentos generales: {all_stats['documentos_generales']}")
        print(f"✅ Total documentos procesados: {all_stats['documentos_procesados']}")
        print(f"✅ Chunks totales generados: {all_stats['chunks_totales']}")
        print(f"\n{'✅' if all_stats['validacion_exitosa'] else '⚠️ '} Validación: {'EXITOSA' if all_stats['validacion_exitosa'] else 'CON ERRORES'}")
        print(f"\n📁 JSONs guardados en: {self.output_path}")
        print("="*60 + "\n")

        return all_stats


def main():
    # Rutas
    data_path = "../data/obras_sociales"
    docs_general_path = "../docs"  # Checklist general del hospital
    output_path = "../data/obras_sociales_json"

    # Convertir
    converter = DocumentToJsonConverter(data_path, output_path)
    stats = converter.process_all_documents(docs_general_path=docs_general_path)

    # Mostrar ejemplos de chunks
    print("\n" + "="*60)
    print("EJEMPLOS DE CHUNKS")
    print("="*60)

    # Ejemplo ENSALUD
    print("\n--- ENSALUD (PDF) ---")
    example_file = Path(output_path) / "ensalud" / "2024-01-04_normativa_chunks.json"
    if example_file.exists():
        with open(example_file, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
            if chunks:
                print(json.dumps(chunks[0], ensure_ascii=False, indent=2))

    # Ejemplo ASI (TABLA con copagos)
    print("\n--- ASI (TABLA DE COPAGOS) ---")
    example_file = Path(output_path) / "asi" / "2024-01-04_normas_chunks.json"
    if example_file.exists():
        with open(example_file, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
            # Buscar chunk de tabla
            for chunk in chunks:
                if chunk.get("es_tabla"):
                    print(json.dumps(chunk, ensure_ascii=False, indent=2))
                    break

    # Ejemplo GENERAL
    print("\n--- CHECKLIST GENERAL HOSPITAL ---")
    example_file = Path(output_path) / "general" / "checklist_general_grupo_pediatrico_chunks.json"
    if example_file.exists():
        with open(example_file, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
            if chunks:
                print(json.dumps(chunks[0], ensure_ascii=False, indent=2))

    return stats


if __name__ == "__main__":
    main()
